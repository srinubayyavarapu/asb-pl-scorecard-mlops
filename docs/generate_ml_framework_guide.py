from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Brand Colors ──
NAVY = RGBColor(0, 51, 102)
DARK_GRAY = RGBColor(50, 50, 50)
MED_GRAY = RGBColor(100, 100, 100)
LIGHT_GRAY = RGBColor(150, 150, 150)
WHITE = RGBColor(255, 255, 255)
ACCENT_BLUE = RGBColor(0, 112, 192)
LIGHT_BG = "E8EEF7"
HEADER_BG = "003366"
ALT_ROW = "F2F6FC"


def shade_cell(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, size=9, color=None, align=None):
    """Set text in a table cell with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Set borders
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
        '</w:tblBorders>'
    )
    tbl_pr.append(borders)

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, HEADER_BG)
        set_cell_text(cell, h, bold=True, size=9, color=WHITE)

    # Data rows with alternating colors
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                shade_cell(cell, ALT_ROW)
            set_cell_text(cell, val, size=9, color=DARK_GRAY)

    # Column widths
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Cm(width)

    doc.add_paragraph("")
    return table


def add_heading_styled(doc, text, level=1):
    """Add a styled heading with navy color."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
        run.font.name = "Calibri"
    return h


def add_body(doc, text):
    """Add body text with consistent styling."""
    p = doc.add_paragraph()
    p.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK_GRAY
    return p


def add_callout(doc, text, bg_color=LIGHT_BG):
    """Add a highlighted callout box."""
    p = doc.add_paragraph()
    p.space_before = Pt(4)
    p.space_after = Pt(8)
    # Shading on paragraph
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>')
    pPr.append(shading)
    # Indentation
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Consolas"
    run.font.color.rgb = NAVY
    return p


def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    p.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
    else:
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
    return p


def add_separator(doc):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="003366"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


# ══════════════════════════════════════════════
#  BUILD DOCUMENT
# ══════════════════════════════════════════════

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.font.color.rgb = DARK_GRAY
style.paragraph_format.space_after = Pt(6)

# ──────────────────────────────────────────────
#  COVER PAGE
# ──────────────────────────────────────────────
for _ in range(4):
    doc.add_paragraph("")

# Top accent line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    '  <w:top w:val="single" w:sz="24" w:space="1" w:color="003366"/>'
    '</w:pBdr>'
)
pPr.append(pBdr)

doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ML Scorecard Framework")
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = NAVY
run.font.name = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Framework & Onboarding Guide")
run.font.size = Pt(16)
run.font.color.rgb = ACCENT_BLUE
run.font.name = "Calibri"

doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    '  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="0070C0"/>'
    '</w:pBdr>'
)
pPr.append(pBdr)

doc.add_paragraph("")
doc.add_paragraph("")

# Cover details table
cover_table = doc.add_table(rows=5, cols=2)
cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cover_data = [
    ("Client", "ASB Bank Limited, New Zealand"),
    ("Engagement", "Credit Risk Scorecard Migration & MLOps"),
    ("Prepared by", "Celebal Technologies (in partnership with Databricks)"),
    ("Version", "1.0"),
    ("Classification", "Confidential"),
]
for i, (label, value) in enumerate(cover_data):
    set_cell_text(cover_table.rows[i].cells[0], label, bold=True, size=10, color=NAVY)
    set_cell_text(cover_table.rows[i].cells[1], value, size=10, color=DARK_GRAY)
    shade_cell(cover_table.rows[i].cells[0], LIGHT_BG)

doc.add_page_break()

# ──────────────────────────────────────────────
#  TABLE OF CONTENTS
# ──────────────────────────────────────────────
add_heading_styled(doc, "Table of Contents", level=1)
add_separator(doc)

toc_items = [
    "1. Overview",
    "2. How the Pipeline Works",
    "3. Project Structure",
    "4. Configuration Reference",
    "    4.1 Data Preparation",
    "    4.2 Feature Engineering",
    "    4.3 Model Training",
    "    4.4 Model Scoring",
    "    4.5 Model Monitoring",
    "5. Running the Pipeline",
    "6. Output Tables",
    "7. Onboarding a New Model (Step-by-Step)",
    "8. Quick Reference",
    "9. Databricks Platform Components",
]
for item in toc_items:
    p = doc.add_paragraph()
    p.space_after = Pt(2)
    run = p.add_run(item)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    if item.startswith("    "):
        p.paragraph_format.left_indent = Cm(1.5)
        run.font.color.rgb = MED_GRAY
    else:
        run.font.color.rgb = NAVY
        run.bold = True

doc.add_page_break()

# ──────────────────────────────────────────────
#  1. OVERVIEW
# ──────────────────────────────────────────────
add_heading_styled(doc, "1. Overview", level=1)
add_separator(doc)

add_body(doc,
    "The ML Scorecard Framework is a configuration-driven pipeline for building, "
    "deploying, and monitoring credit risk scorecards on the Databricks platform."
)
add_body(doc,
    "The framework separates what you want to do (configuration) from how it is "
    "done (code). You describe your model's data sources, features, algorithms, "
    "scoring rules, and monitoring thresholds in simple YAML configuration files. "
    "The framework's generic notebooks read those configurations and execute the "
    "full ML lifecycle automatically."
)

add_body(doc, "This means:")
add_bullet(doc, "No code changes are needed when onboarding a new scorecard model")
add_bullet(doc, "All models follow the same tested, validated pipeline")
add_bullet(doc, "Configuration changes are versioned in Git alongside the code")
add_bullet(doc, "Each model's pipeline is independently runnable and monitorable")

doc.add_paragraph("")
add_callout(doc, "Key Principle:  One pipeline. Many models. No code changes.")

# ──────────────────────────────────────────────
#  2. HOW THE PIPELINE WORKS
# ──────────────────────────────────────────────
doc.add_page_break()
add_heading_styled(doc, "2. How the Pipeline Works", level=1)
add_separator(doc)

add_body(doc,
    "The framework executes 5 stages in sequence. Each stage is a separate "
    "Databricks notebook that reads your model's configuration and performs "
    "its designated task. In the Databricks Workflows UI, you can see each "
    "stage as an individual task with its own status, duration, and logs."
)

add_styled_table(doc,
    ["Stage", "Notebook", "What It Does"],
    [
        ["1. Data Preparation", "01_data_prep.py",
         "Reads source tables from the Silver layer, joins them as specified in config, "
         "applies exclusion rules, defines good/bad outcomes, and creates stratified "
         "Dev / Holdout / Out-of-Time population splits. Writes ML-ready data to the Gold layer."],
        ["2. Feature Engineering", "02_feature_engineering.py",
         "Calculates Weight of Evidence (WoE) and Information Value (IV) for each feature. "
         "Selects features that meet the IV threshold. Saves banding lookup tables and the "
         "feature store for downstream use."],
        ["3. Model Training", "03_model_training.py",
         "Trains the champion model (Logistic Regression) and challenger models "
         "(Random Forest, Neural Network) with hyperparameter grid search. Logs all parameters, "
         "metrics, and model artifacts to MLflow. Registers the best model in Unity Catalog."],
        ["4. Model Scoring", "04_model_scoring.py",
         "Loads the registered model, scores all records, applies scorecard scaling "
         "(Points-to-Double-the-Odds), assigns risk grades (A1 through D), and writes "
         "the scored output to the Gold layer."],
        ["5. Model Monitoring", "05_model_monitoring.py",
         "Calculates Gini, KS, and AUC for each population (Dev, Holdout, Out-of-Time). "
         "Computes PSI to detect distribution drift. Raises warning or critical alerts "
         "if any metric breaches configured thresholds."],
    ],
    col_widths=[3.5, 4, 10]
)

add_body(doc, "Task Dependency Chain:")
add_callout(doc,
    "data_prep  -->  feature_engineering  -->  model_training  -->  model_scoring  -->  model_monitoring"
)

# ──────────────────────────────────────────────
#  3. PROJECT STRUCTURE
# ──────────────────────────────────────────────
doc.add_page_break()
add_heading_styled(doc, "3. Project Structure", level=1)
add_separator(doc)

add_body(doc,
    "The framework follows a clean separation between configuration and code. "
    "Configuration files describe what each model needs. Notebook files contain "
    "the generic logic that all models share."
)

add_heading_styled(doc, "Configuration Files", level=2)
add_styled_table(doc,
    ["File", "Purpose"],
    [
        ["configs/ml/model_registry.yaml",
         "Master registry that lists all models. Each entry points to a config folder."],
        ["configs/ml/<model_name>/data_prep.yaml",
         "Data sources, join logic, good/bad definition, exclusions, sampling ratios"],
        ["configs/ml/<model_name>/features.yaml",
         "Feature list, types, banding methods, IV thresholds, feature store location"],
        ["configs/ml/<model_name>/training.yaml",
         "MLflow experiment name, algorithms, hyperparameters, validation thresholds"],
        ["configs/ml/<model_name>/scoring.yaml",
         "Scorecard scaling parameters, risk grade definitions, output table"],
        ["configs/ml/<model_name>/monitoring.yaml",
         "Metrics to track, drift thresholds, alerting configuration"],
    ],
    col_widths=[6.5, 11]
)

add_heading_styled(doc, "Notebook Files (Generic - Never Modified)", level=2)
add_styled_table(doc,
    ["File", "Role"],
    [
        ["notebooks/ml/00_ml_config_loader.py", "Shared utility that loads and resolves configs"],
        ["notebooks/ml/01_data_prep.py", "Stage 1: Data preparation"],
        ["notebooks/ml/02_feature_engineering.py", "Stage 2: WoE/IV analysis and feature selection"],
        ["notebooks/ml/03_model_training.py", "Stage 3: Model training with MLflow tracking"],
        ["notebooks/ml/04_model_scoring.py", "Stage 4: Batch scoring and risk grading"],
        ["notebooks/ml/05_model_monitoring.py", "Stage 5: Performance monitoring and drift detection"],
    ],
    col_widths=[7, 10.5]
)

add_heading_styled(doc, "Job Definition", level=2)
add_styled_table(doc,
    ["File", "Purpose"],
    [
        ["resources/ml_scorecard_job.yml",
         "Databricks Asset Bundle job definition. Defines the 5-task chain with dependencies "
         "for each model. One job block per model."],
    ],
    col_widths=[6.5, 11]
)

# ──────────────────────────────────────────────
#  4. CONFIGURATION REFERENCE
# ──────────────────────────────────────────────
doc.add_page_break()
add_heading_styled(doc, "4. Configuration Reference", level=1)
add_separator(doc)

add_body(doc,
    "Each model has 5 YAML configuration files. Below is a reference of every "
    "setting available in each file, what it controls, and an example value."
)

# 4.1
add_heading_styled(doc, "4.1  data_prep.yaml - Data Preparation", level=2)
add_styled_table(doc,
    ["Setting", "Description", "Example Value"],
    [
        ["source_tables", "Named map of Silver tables to read", "base: catalog.silver.hlacctbase_final"],
        ["joins", "Ordered list of join operations (left, right, on, how)", "left join gdw_extract on account_key"],
        ["good_bad_definition.target_column", "Column containing the default flag", "default_flag"],
        ["good_bad_definition.bad_condition", "SQL condition that defines a 'bad' outcome", "default_flag = 1"],
        ["good_bad_definition.exclusions", "List of SQL conditions for records to exclude", "account_status = 'CLOSED'"],
        ["sampling.dev_ratio", "Proportion of in-time data for development", "0.70"],
        ["sampling.holdout_ratio", "Proportion of in-time data for validation", "0.30"],
        ["sampling.oot_start_date", "Date boundary for Out-of-Time population", "2023-07-01"],
        ["sampling.stratify_by", "Column to stratify the split on", "default_flag"],
        ["sampling.random_seed", "Random seed for reproducibility", "42"],
        ["output_table", "Fully qualified table name for the output", "catalog.gold.hl_scorecard_dev_data"],
    ],
    col_widths=[5, 7, 5.5]
)

# 4.2
add_heading_styled(doc, "4.2  features.yaml - Feature Engineering", level=2)
add_styled_table(doc,
    ["Setting", "Description", "Example Value"],
    [
        ["features[].name", "Feature column name", "credit_score"],
        ["features[].type", "continuous or categorical", "continuous"],
        ["features[].banding_method", "auto (quantile bins) or manual (defined bands)", "auto"],
        ["features[].monotonicity", "Expected WoE direction: ascending, descending, or null", "descending"],
        ["iv_thresholds.exclude_below", "IV below this value -> feature excluded", "0.02"],
        ["iv_thresholds.suspicious_above", "IV above this value -> investigate for leakage", "0.50"],
        ["auto_banding.max_bins", "Maximum number of bins for auto banding", "10"],
        ["auto_banding.min_bin_size", "Minimum population percentage per bin", "0.05"],
        ["woe_iv_table", "Where to save WoE/IV analysis results", "catalog.gold.hl_woe_iv"],
        ["feature_store_table", "Where to save the final feature store", "catalog.ml.hl_feature_store"],
    ],
    col_widths=[5, 7, 5.5]
)

# 4.3
add_heading_styled(doc, "4.3  training.yaml - Model Training", level=2)
add_styled_table(doc,
    ["Setting", "Description", "Example Value"],
    [
        ["mlflow.experiment_name", "MLflow experiment path in workspace", "/Shared/ml/hl_scorecard_experiments"],
        ["mlflow.model_name", "Unity Catalog model registry name", "catalog.ml.hl_behaviour_scorecard"],
        ["algorithms.<name>.class", "Scikit-learn class name", "LogisticRegression"],
        ["algorithms.<name>.role", "champion or challenger", "champion"],
        ["algorithms.<name>.params", "Fixed hyperparameters", "penalty: l2, solver: lbfgs"],
        ["algorithms.<name>.grid_search", "Parameters to search over", "C: [0.01, 0.1, 1.0, 10.0]"],
        ["validation_thresholds.gini.min", "Minimum Gini coefficient required", "0.40"],
        ["validation_thresholds.ks.min", "Minimum KS statistic required", "0.25"],
        ["validation_thresholds.auc.min", "Minimum AUC-ROC required", "0.70"],
        ["validation_thresholds.psi.green", "PSI below this = stable", "0.10"],
    ],
    col_widths=[5, 7, 5.5]
)

# 4.4
add_heading_styled(doc, "4.4  scoring.yaml - Model Scoring", level=2)
add_styled_table(doc,
    ["Setting", "Description", "Example Value"],
    [
        ["scorecard_scaling.target_score", "Credit score at the target odds ratio", "600"],
        ["scorecard_scaling.target_odds", "Good:Bad ratio at the target score", "50 (50:1)"],
        ["scorecard_scaling.pdo", "Points to Double the Odds", "20"],
        ["risk_grades[].grade", "Risk grade label", "A1"],
        ["risk_grades[].score_min / score_max", "Score range for this grade", "750 - 999"],
        ["risk_grades[].pd_min / pd_max", "PD range for this grade", "0.000 - 0.005"],
        ["output_table", "Where to write scored output", "catalog.gold.hl_scored_output"],
    ],
    col_widths=[5, 7, 5.5]
)

# 4.5
add_heading_styled(doc, "4.5  monitoring.yaml - Model Monitoring", level=2)
add_styled_table(doc,
    ["Setting", "Description", "Example Value"],
    [
        ["metrics", "List of metrics to track over time", "gini, ks, auc, psi"],
        ["drift_thresholds.psi.warning", "PSI value that triggers a warning", "0.10"],
        ["drift_thresholds.psi.critical", "PSI value that triggers a critical alert", "0.25"],
        ["drift_thresholds.gini_drop.warning", "Gini drop from baseline that triggers warning", "0.05"],
        ["monitoring_table", "Where to append monitoring results", "catalog.ml.hl_monitoring_log"],
        ["baseline_table", "Where to store the initial baseline metrics", "catalog.ml.hl_monitoring_baseline"],
    ],
    col_widths=[5, 7, 5.5]
)

# ──────────────────────────────────────────────
#  5. RUNNING THE PIPELINE
# ──────────────────────────────────────────────
doc.add_page_break()
add_heading_styled(doc, "5. Running the Pipeline", level=1)
add_separator(doc)

add_heading_styled(doc, "Option A: Databricks CLI (Recommended)", level=2)
add_body(doc, "Deploy the bundle and trigger the job from your terminal:")
add_callout(doc, "databricks bundle deploy --target dev")
add_callout(doc, "databricks bundle run ml_hl_behaviour_scorecard --target dev")

add_heading_styled(doc, "Option B: Databricks Workflows UI", level=2)
add_body(doc,
    "Navigate to Workflows > Jobs in the Databricks workspace. "
    "Select the job (e.g., ASB ML Scorecard Pipeline) and click Run Now. "
    "You can monitor each task's progress, view logs, and check outputs."
)

add_heading_styled(doc, "Option C: Run Notebooks Manually", level=2)
add_body(doc,
    "Open each notebook in order (01 through 05) in the Databricks workspace. "
    "Set two widget parameters at the top of each notebook:"
)
add_bullet(doc, " = your model name from the registry (e.g., hl_behaviour_scorecard)", bold_prefix="model_name")
add_bullet(doc, " = your Unity Catalog name (e.g., asb_dev)", bold_prefix="catalog")

# ──────────────────────────────────────────────
#  6. OUTPUT TABLES
# ──────────────────────────────────────────────
add_heading_styled(doc, "6. Output Tables", level=1)
add_separator(doc)

add_body(doc,
    "After a successful pipeline run, the following tables are created or updated:"
)

add_styled_table(doc,
    ["Table Name", "Schema", "Description"],
    [
        ["hl_scorecard_dev_data", "retail_gold",
         "ML-ready dataset with Dev, Holdout, and OOT population labels. "
         "Contains all features and the target variable after exclusions and sampling."],
        ["hl_woe_iv", "retail_gold",
         "Weight of Evidence and Information Value analysis results. "
         "One row per feature per bin, with WoE and IV values."],
        ["hl_feature_store", "retail_ml",
         "Feature store containing the selected features and metadata. "
         "Includes the _selected_features column listing which features were chosen."],
        ["hl_scored_output", "retail_gold",
         "Every record scored with PD estimate, credit score (0-999), "
         "and risk grade (A1 through D). Ready for downstream reporting."],
        ["hl_monitoring_log", "retail_ml",
         "Monitoring metrics (Gini, KS, AUC, PSI) per population per run. "
         "Includes alert level (none, warning, critical) per metric."],
        ["hl_model_load_uri", "retail_ml",
         "Stores the URI of the best model from the latest training run. "
         "Used by the scoring notebook to load the correct model."],
    ],
    col_widths=[4.5, 3, 10]
)

# ──────────────────────────────────────────────
#  7. ONBOARDING A NEW MODEL
# ──────────────────────────────────────────────
doc.add_page_break()
add_heading_styled(doc, "7. Onboarding a New Model", level=1)
add_separator(doc)

add_body(doc,
    "Adding a new scorecard model to the framework requires no code changes. "
    "You create configuration files, register the model, and add a job definition. "
    "The same 5 generic notebooks handle every model."
)

add_body(doc, "Follow these 5 steps:")

doc.add_paragraph("")

# Step 1
add_heading_styled(doc, "Step 1: Create a Configuration Folder", level=2)
add_body(doc,
    "Create a new folder under configs/ml/ named after your model. "
    "Copy the 5 YAML files from an existing model's folder as a starting template."
)
add_callout(doc,
    "Example:\n"
    "  configs/ml/cc_behaviour_scorecard/\n"
    "      data_prep.yaml\n"
    "      features.yaml\n"
    "      training.yaml\n"
    "      scoring.yaml\n"
    "      monitoring.yaml"
)

# Step 2
add_heading_styled(doc, "Step 2: Edit the Configuration Files", level=2)
add_body(doc,
    "Open each YAML file and update the values for your new model. "
    "The table below shows what typically needs to change in each file:"
)

add_styled_table(doc,
    ["Config File", "What to Change"],
    [
        ["data_prep.yaml",
         "Source table names pointing to your model's Silver data. "
         "Join logic if your model needs different table combinations. "
         "Exclusion rules specific to this product. Output table name."],
        ["features.yaml",
         "Feature names and types for your model. "
         "Banding methods (auto vs manual) per feature. "
         "WoE/IV and feature store output table names."],
        ["training.yaml",
         "MLflow experiment name (unique per model). "
         "Model registry name in Unity Catalog. "
         "Algorithm hyperparameters if different from defaults."],
        ["scoring.yaml",
         "Scored output table name. "
         "Risk grade definitions if your product uses different grade bands."],
        ["monitoring.yaml",
         "Monitoring log table name. "
         "Drift thresholds if your model has different tolerance levels."],
    ],
    col_widths=[4, 13.5]
)

# Step 3
add_heading_styled(doc, "Step 3: Register in the Model Registry", level=2)
add_body(doc,
    "Open configs/ml/model_registry.yaml and add a new entry for your model:"
)
add_callout(doc,
    "cc_behaviour_scorecard:\n"
    "    description: \"Credit Card Behaviour Scorecard\"\n"
    "    product: \"credit_card\"\n"
    "    use_case: \"credit_risk_pd\"\n"
    "    target_variable: \"default_flag\"\n"
    "    primary_key: \"account_key\"\n"
    "    champion_algorithm: \"logistic_regression\"\n"
    "    challenger_algorithms: [\"random_forest\"]\n"
    "    config_dir: \"configs/ml/cc_behaviour_scorecard\"\n"
    "    is_active: true"
)

# Step 4
add_heading_styled(doc, "Step 4: Add a Databricks Job", level=2)
add_body(doc,
    "Open resources/ml_scorecard_job.yml and add a new job block for your model. "
    "Copy the existing job block and change only two values:"
)

add_styled_table(doc,
    ["Field to Change", "What to Set"],
    [
        ["Job name", "A unique name for this model's job (e.g., asb-ml-cc-scorecard-dev)"],
        ["model_name parameter default", "Your model's key from the registry (e.g., cc_behaviour_scorecard)"],
    ],
    col_widths=[5, 12.5]
)

add_body(doc,
    "The 5 task definitions (data_prep through model_monitoring) remain identical. "
    "They reference the same generic notebooks. Only the job name and the default "
    "model_name parameter value change."
)

# Step 5
add_heading_styled(doc, "Step 5: Deploy and Run", level=2)
add_body(doc, "Deploy the updated bundle to your Databricks workspace and trigger the new job:")
add_callout(doc,
    "databricks bundle deploy --target dev\n"
    "databricks bundle run ml_cc_behaviour_scorecard --target dev"
)

add_body(doc,
    "The pipeline will execute all 5 stages using your new model's configuration. "
    "Monitor progress in the Databricks Workflows UI."
)

# ──────────────────────────────────────────────
#  8. QUICK REFERENCE
# ──────────────────────────────────────────────
doc.add_page_break()
add_heading_styled(doc, "8. Quick Reference", level=1)
add_separator(doc)

add_body(doc,
    "The table below summarises what needs to change for common scenarios. "
    "The key takeaway: notebooks are never modified."
)

add_styled_table(doc,
    ["Scenario", "Change Configs?", "Change Notebooks?", "Change Job?"],
    [
        ["Onboard a new model", "Yes - new folder with 5 YAMLs", "No", "Yes - new job block"],
        ["Add or remove features", "Yes - features.yaml", "No", "No"],
        ["Change algorithm hyperparameters", "Yes - training.yaml", "No", "No"],
        ["Change risk grade definitions", "Yes - scoring.yaml", "No", "No"],
        ["Change drift alert thresholds", "Yes - monitoring.yaml", "No", "No"],
        ["Change source data tables", "Yes - data_prep.yaml", "No", "No"],
        ["Change Dev/Holdout split ratio", "Yes - data_prep.yaml", "No", "No"],
        ["Add a new algorithm", "Yes - training.yaml", "No", "No"],
    ],
    col_widths=[5, 5, 3.5, 4]
)

doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Notebooks never change.  Configuration always changes.")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = NAVY

# ──────────────────────────────────────────────
#  9. PLATFORM COMPONENTS
# ──────────────────────────────────────────────
doc.add_paragraph("")
add_heading_styled(doc, "9. Databricks Platform Components", level=1)
add_separator(doc)

add_body(doc,
    "The framework is built entirely on Databricks-native services. "
    "No external tools or custom infrastructure is required."
)

add_styled_table(doc,
    ["Component", "Databricks Service", "How It Is Used"],
    [
        ["Data Storage", "Delta Lake + Unity Catalog",
         "All tables stored in Bronze, Silver, Gold, and ML schemas"],
        ["Experiment Tracking", "MLflow",
         "Parameters, metrics, and model artifacts logged per training run"],
        ["Model Registry", "Unity Catalog Model Registry",
         "Model versioning with @Champion and @Challenger aliases"],
        ["Compute", "Serverless",
         "No cluster management required. Auto-scales per task."],
        ["Orchestration", "Databricks Workflows",
         "Multi-task jobs with dependencies. One job per model."],
        ["Governance", "Unity Catalog",
         "Access control, data lineage, and audit trail for regulatory compliance"],
        ["Deployment", "Databricks Asset Bundles (DAB)",
         "Declarative YAML-based deployment across Dev, Staging, and Production"],
    ],
    col_widths=[3.5, 5, 9]
)

# ──────────────────────────────────────────────
#  SAVE
# ──────────────────────────────────────────────
output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, "ML_Scorecard_Framework_Guide.docx")
doc.save(output_path)
print(f"Document saved: {output_path}")
