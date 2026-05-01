"""Extend ASB Credit Risk Platform Design Document V1 with CT's Catalog Structure deep-dive.

Final iteration with explicit schema and environment explanations:

Sections inserted under existing 'Catalog Structure' Heading 1:
  - Catalog Design Principles (H2) + table
  - Schema Pattern Inside a Catalog (H2) + 3 explanation tables
      (medallion + use-case schemas; concrete use-case schemas per team; environment roles)
  - Catalog Design Options (H2) + intro
      - Option A - Single Catalog per Environment (H3) + structure table + pros/cons
      - Option B - Catalog per Team per Environment (H3) + structure table + pros/cons
  - Comparison Across Design Principles (H2) + table
  - Recommendation (H2) + bullets
  - Naming Conventions (H2) + table
  - Permission Model (H2) + table mapping the 29 FTE

Removes any prior CT extension content before inserting.
"""
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W_T = f"{{{W_NS}}}t"
W_P_STYLE = f"{{{W_NS}}}pStyle"
W_VAL = f"{{{W_NS}}}val"

DOCS = Path(
    r"C:\Users\SrinuBayyavarapu\OneDrive - Celebal Technologies Private Limited"
    r"\ASB-Bank-SAS-to-Dbx-Migration\docs"
)
SRC = DOCS / "ASB Credit Risk Platform Design Document_V1.docx"
BACKUP = DOCS / "ASB Credit Risk Platform Design Document_V1.backup.docx"

shutil.copy(SRC, BACKUP)
print(f"Backup -> {BACKUP.name}")

doc = Document(str(SRC))
body = doc.element.body


def first_text(elem):
    t = elem.find(f".//{W_T}")
    return t.text if t is not None else ""


def is_table_with_first_cell(elem, expected):
    return elem.tag.endswith("}tbl") and first_text(elem).strip() == expected


def is_heading_para(elem, heading_style, expected_text):
    if not elem.tag.endswith("}p"):
        return False
    pStyle = elem.find(f".//{W_P_STYLE}")
    if pStyle is None or pStyle.get(W_VAL) != heading_style:
        return False
    full = "".join(t.text for t in elem.iter(W_T) if t.text)
    return full.strip() == expected_text


# ===================================================================
# STEP 1: Find anchors and remove any prior CT extension content
# ===================================================================
anchor_idx = None
key_tech_idx = None
for i, child in enumerate(body):
    if anchor_idx is None and is_table_with_first_cell(child, "Dev_credit_risk"):
        anchor_idx = i
    if anchor_idx is not None and is_heading_para(child, "Heading1", "Key Technologies"):
        key_tech_idx = i
        break

assert anchor_idx is not None and key_tech_idx is not None

removed = 0
for elem in list(body[anchor_idx + 1 : key_tech_idx]):
    body.remove(elem)
    removed += 1
print(f"Removed {removed} prior CT extension elements")

target = None
for p in doc.paragraphs:
    if p.text.strip() == "Key Technologies" and p.style.name == "Heading 1":
        target = p
        break
assert target is not None


# ===================================================================
# Helpers
# ===================================================================
def set_table_full_width(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)
    table.autofit = True
    table.allow_autofit = True


def add_para(text, style="Normal"):
    new_p = doc.add_paragraph(text, style=style)
    target._element.addprevious(new_p._element)
    return new_p


def add_heading(text, level):
    return add_para(text, style=f"Heading {level}")


def add_bullet(text):
    return add_para(f"•  {text}", style="Normal")


def write_cell(cell, val):
    """val may be:
        - str: single-line cell
        - list[str]: multi-paragraph cell
        - tuple(header_str, list[str]): bold header followed by multi-paragraph content
    """
    if isinstance(val, tuple):
        header, lines = val
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        for line in lines:
            cell.add_paragraph(line)
    elif isinstance(val, list):
        cell.text = val[0] if val else ""
        for line in val[1:]:
            cell.add_paragraph(line)
    else:
        cell.text = val


def add_table(headers, rows_data, style="Grid Table 5 Dark Accent 2"):
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    try:
        table.style = style
    except KeyError:
        table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for r, row in enumerate(rows_data):
        for c, val in enumerate(row):
            write_cell(table.cell(r + 1, c), val)
    set_table_full_width(table)
    target._element.addprevious(table._element)
    return table


# Schema content per team
RM_SCHEMAS = ["bronze", "silver", "gold", "pl_scorecard", "cc_scorecard"]
RI_SCHEMAS = ["bronze", "silver", "gold", "reporting"]
CA_SCHEMAS = ["bronze", "silver", "gold", "model_dev"]
FR_SCHEMAS = ["bronze", "silver", "gold", "market_risk", "climate_risk"]

# Option A: schemas under a single shared catalog carry the team prefix
RM_PREFIXED = [f"retail_modelling_{s}" for s in RM_SCHEMAS]
RI_PREFIXED = [f"retail_insights_{s}" for s in RI_SCHEMAS]
CA_PREFIXED = [f"corporate_analytics_{s}" for s in CA_SCHEMAS]
FR_PREFIXED = [f"financial_risk_{s}" for s in FR_SCHEMAS]


# ===================================================================
# 1. CATALOG DESIGN PRINCIPLES
# ===================================================================
add_heading("Catalog Design Principles", 2)
add_para(
    "The catalog topology serves the Risk Division - 29 FTE across four modelling "
    "teams plus the Model Risk and Assurance validator team. The design is judged "
    "against six principles drawn from RBNZ BS11 and APRA CPS 234 obligations and "
    "Unity Catalog design fundamentals."
)
add_table(
    headers=["#", "Principle", "What it requires"],
    rows_data=[
        ["1", "Environment isolation",
         "Dev, staging, and production data are physically separable; non-prod work cannot reach prod data."],
        ["2", "Team ownership",
         "Each Risk Division team has its own data ownership boundary it controls grants on."],
        ["3", "Regulatory separation of duties",
         "Production writes restricted to service principals; humans never deploy directly."],
        ["4", "Validator scoping",
         "The Model Risk and Assurance team (2 FTE) reads prod registry, monitoring, and lineage with no write privilege and no compute."],
        ["5", "Bounded blast radius",
         "A failure in one team's catalog cannot affect another team's data."],
        ["6", "Scale across the 29-FTE roadmap",
         "Onboarding a new team is pattern repetition, not architectural change."],
    ],
)

# ===================================================================
# 2. SCHEMA PATTERN INSIDE A CATALOG  (NEW - foundation section)
# ===================================================================
add_heading("Schema Pattern Inside a Catalog", 2)
add_para(
    "Every catalog in the design - whether shared across teams (Option A below) or "
    "owned by a single team (Option B below) - holds the same kinds of schemas "
    "internally. There are two kinds: the three medallion layers, and one or more "
    "use-case schemas. The table below explains what each kind contains."
)
add_table(
    headers=["Schema", "Kind", "Purpose"],
    rows_data=[
        [
            "bronze",
            "Medallion - raw",
            "Raw data ingested from Snowflake or external files. Untransformed, append-only, fully auditable. The system of record for what arrived from source.",
        ],
        [
            "silver",
            "Medallion - cleansed",
            "Validated, deduplicated, type-cast data with SCD Type 2 history applied to dimensions. The trusted layer that all downstream consumers read from.",
        ],
        [
            "gold",
            "Medallion - model-ready",
            "Feature-engineered datasets joining silver tables and applying business logic for a specific use case. Rebuilt deterministically on every pipeline run, guaranteeing point-in-time consistency.",
        ],
        [
            "<usecase>",
            "Use case",
            "All ML assets for one specific model or reporting product live in this schema together - feature tables, registered model, inference outputs, monitoring tables, evaluation results. The use case is fully addressable by one fully qualified name (e.g., prod_retail_modelling.pl_scorecard).",
        ],
    ],
)
add_para(
    "Across the four Risk Division modelling teams, the use-case schemas planned for "
    "the target state are listed below. The medallion schemas (bronze, silver, gold) "
    "are universal and required for every catalog; the use-case schemas are specific "
    "to the work the owning team performs."
)
add_table(
    headers=["Use-case schema", "Owning team", "What it holds"],
    rows_data=[
        [
            "pl_scorecard",
            "Retail Modelling",
            "Personal Loan Application Scorecard - feature store, registered model, scored output, monitoring metrics.",
        ],
        [
            "cc_scorecard",
            "Retail Modelling",
            "Credit Card Behaviour Scorecard - same kinds of ML assets as pl_scorecard.",
        ],
        [
            "reporting",
            "Retail Insights",
            "Reporting outputs and downstream BI feeds.",
        ],
        [
            "model_dev",
            "Corporate Analytics",
            "Models the team is currently developing.",
        ],
        [
            "market_risk",
            "Financial Risk",
            "Market Risk model assets.",
        ],
        [
            "climate_risk",
            "Financial Risk",
            "Climate Risk model assets.",
        ],
    ],
)
add_para(
    "Each catalog is replicated across three environments. The environments serve "
    "specific roles in the development lifecycle, and only certain principals can "
    "write to each."
)
add_table(
    headers=["Environment", "Purpose", "Who can write"],
    rows_data=[
        [
            "dev",
            "Development sandbox where data scientists prototype, experiment, and iterate on new models.",
            "Data scientists have read-write access. The CI service principal also writes here for build artefacts.",
        ],
        [
            "stg",
            "Pre-production staging where CI integration tests run end-to-end before code is promoted to production.",
            "The CI service principal has write access. Data scientists have read-only access.",
        ],
        [
            "prod",
            "Live production where the champion model scores real customer data on a regular cadence.",
            "Only the CD service principal can write here. Data scientists have read-only access for incident response.",
        ],
    ],
)

# ===================================================================
# 3. CATALOG DESIGN OPTIONS
# ===================================================================
add_heading("Catalog Design Options", 2)
add_para(
    "Two catalog topologies are evaluated against the principles, using the actual "
    "Risk Division team names as the basis for catalog naming. Both options use the "
    "schema pattern described above and are replicated across the three environments "
    "(dev, stg, prod)."
)

# --- Option A ---
add_heading("Option A - Single Catalog per Environment", 3)
add_para(
    "A single catalog per environment shared by all four Risk Division modelling "
    "teams. Inside that one catalog, every team's schemas carry a team prefix to keep "
    "namespaces visually separated. For example, the Retail Modelling team's bronze, "
    "silver, gold, pl_scorecard, and cc_scorecard schemas appear as "
    "retail_modelling_bronze, retail_modelling_silver, retail_modelling_gold, "
    "retail_modelling_pl_scorecard, and retail_modelling_cc_scorecard. The same set "
    "of schemas exists in risk_dev, risk_stg, and risk_prod."
)
add_table(
    headers=["Owning team", "Schemas in risk_dev", "Schemas in risk_stg", "Schemas in risk_prod"],
    rows_data=[
        ["Retail Modelling",     RM_PREFIXED, RM_PREFIXED, RM_PREFIXED],
        ["Retail Insights",      RI_PREFIXED, RI_PREFIXED, RI_PREFIXED],
        ["Corporate Analytics",  CA_PREFIXED, CA_PREFIXED, CA_PREFIXED],
        ["Financial Risk",       FR_PREFIXED, FR_PREFIXED, FR_PREFIXED],
    ],
)
add_para(
    "Pros: fewest catalogs to provision; cross-team lineage is in scope by default; "
    "single operational model across the Risk Division."
)
add_para(
    "Cons: a failure in one team's schemas can affect any other team's schemas in the "
    "same catalog; validator scoping requires schema-level grants on every schema "
    "rather than a single catalog grant; team ownership is implicit, enforced only by "
    "naming convention."
)

# --- Option B ---
add_heading("Option B - Catalog per Team per Environment", 3)
add_para(
    "Each Risk Division team owns its own catalog, and that catalog is replicated "
    "across all three environments. Catalog names follow the pattern <env>_<team>, "
    "the same convention as ASB's existing workspace naming. Inside every catalog, "
    "the structure is the medallion + use-case pattern described above: a bronze "
    "schema, a silver schema, a gold schema, plus one use-case schema per model or "
    "reporting product the team owns."
)
add_para(
    "The table below shows the full catalog and schema layout for all four Risk "
    "Division modelling teams across dev, staging, and production. Each row is one "
    "team; each column is one environment. The first line in each cell (in bold) is "
    "the catalog name; the indented lines beneath it are the schemas inside that "
    "catalog."
)
add_table(
    headers=["Team", "Dev Catalog", "Staging Catalog", "Production Catalog"],
    rows_data=[
        [
            "Retail Modelling (10 FTE)",
            ("dev_retail_modelling", [f"  {s}" for s in RM_SCHEMAS]),
            ("stg_retail_modelling", [f"  {s}" for s in RM_SCHEMAS]),
            ("prod_retail_modelling", [f"  {s}" for s in RM_SCHEMAS]),
        ],
        [
            "Retail Insights (9 FTE)",
            ("dev_retail_insights", [f"  {s}" for s in RI_SCHEMAS]),
            ("stg_retail_insights", [f"  {s}" for s in RI_SCHEMAS]),
            ("prod_retail_insights", [f"  {s}" for s in RI_SCHEMAS]),
        ],
        [
            "Corporate Analytics (6 FTE)",
            ("dev_corporate_analytics", [f"  {s}" for s in CA_SCHEMAS]),
            ("stg_corporate_analytics", [f"  {s}" for s in CA_SCHEMAS]),
            ("prod_corporate_analytics", [f"  {s}" for s in CA_SCHEMAS]),
        ],
        [
            "Financial Risk (2 FTE)",
            ("dev_financial_risk", [f"  {s}" for s in FR_SCHEMAS]),
            ("stg_financial_risk", [f"  {s}" for s in FR_SCHEMAS]),
            ("prod_financial_risk", [f"  {s}" for s in FR_SCHEMAS]),
        ],
    ],
)
add_para(
    "Pros: each team has a clean ownership boundary; validator access is a single "
    "grant per catalog per environment; blast radius is contained to one team; "
    "pattern repeats cleanly as remaining SAS models migrate; aligns with ASB's "
    "existing workspace naming."
)
add_para(
    "Cons: more catalogs to provision (4 teams x 3 envs = 12 catalogs); cross-team "
    "data sharing requires explicit grants between catalogs."
)

# ===================================================================
# 4. COMPARISON
# ===================================================================
add_heading("Comparison Across Design Principles", 2)
add_para(
    "The two options scored against the principles. Double tick = strong fit, "
    "single tick = adequate fit, cross = weak fit."
)
add_table(
    headers=["Principle", "A: Single Catalog per Env", "B: Catalog per Team per Env"],
    rows_data=[
        ["Environment isolation",                "✓✓", "✓✓"],
        ["Team ownership",                       "✗",       "✓✓"],
        ["Regulatory separation of duties",      "✓",       "✓✓"],
        ["Validator scoping",                    "✗",       "✓✓"],
        ["Bounded blast radius",                 "✗",       "✓✓"],
        ["Scale across the Risk Division",       "✓",       "✓✓"],
        ["Operational simplicity",               "✓✓", "✓"],
    ],
)

# ===================================================================
# 5. RECOMMENDATION (bullets)
# ===================================================================
add_heading("Recommendation", 2)
add_para(
    "Option B - Catalog per Team per Environment - is the recommended target state "
    "for the Risk Division. The Retail Modelling team (10 FTE) is onboarded first via "
    "the MVP, with dev_retail_modelling, stg_retail_modelling, and prod_retail_modelling "
    "as the first three catalogs created. The same pattern applies to the remaining "
    "three modelling teams as their SAS models migrate."
)
add_para("The recommendation rests on the following points:")
add_bullet(
    "Aligns with ASB's existing workspace convention; the catalog name is the "
    "workspace name and no new mental model is introduced for the team."
)
add_bullet(
    "Regulatory separation of duties is enforced at the catalog level; the Model "
    "Risk and Assurance team gets a single grant per team-catalog they audit, not a "
    "per-schema grant matrix."
)
add_bullet(
    "Blast radius is contained to one team; a failed pipeline or bad DDL in one "
    "team's catalog cannot affect another team's data."
)
add_bullet(
    "Scales cleanly across the Risk Division: each new team adds three catalogs "
    "(dev, stg, prod) with no architectural change and no re-platforming."
)
add_bullet(
    "Pattern repeats identically for future migrations; the next team is onboarded "
    "by provisioning three catalogs with the same medallion + use-case schema layout."
)
add_para(
    "Option A is rejected: shared blast radius and schema-level validator scoping "
    "fail the regulator-facing principles - team ownership, validator scoping, and "
    "blast radius."
)

# ===================================================================
# 6. NAMING CONVENTIONS
# ===================================================================
add_heading("Naming Conventions", 2)
add_para(
    "The conventions below apply to the Option B target state. Catalog, schema, and "
    "table names form a coherent hierarchy aligned with ASB's workspace naming. "
    "Models, feature tables, monitoring tables, and inference tables for a single use "
    "case live together in a single use-case schema, so the use case is fully "
    "addressable via one fully qualified name."
)
add_table(
    headers=["Asset", "Pattern", "Example"],
    rows_data=[
        ["Catalog",            "<env>_<team>",                         "prod_retail_modelling"],
        ["Bronze schema",      "bronze",                               "prod_retail_modelling.bronze"],
        ["Silver schema",      "silver",                               "prod_retail_modelling.silver"],
        ["Gold schema",        "gold",                                 "prod_retail_modelling.gold"],
        ["Use-case schema",    "<usecase>",                            "prod_retail_modelling.pl_scorecard"],
        ["Bronze table",       "<source_system>_<entity>",             "bronze.snowflake_applications"],
        ["Silver table",       "<entity>",                             "silver.applications"],
        ["Gold table",         "<usecase>_<entity>",                   "gold.pl_scorecard_data"],
        ["Use-case table",     "<usecase>_<asset>",                    "pl_scorecard.feature_store"],
        ["Registered model",   "<usecase>_<model_name>",               "pl_scorecard.application_scorecard"],
        ["MLflow experiment",  "/Shared/ml/<usecase>/<model_name>",    "/Shared/ml/pl_scorecard/application_scorecard"],
        ["Job",                "asb-<pipeline>-<usecase>-<env>",       "asb-ml-pl-training-prod"],
    ],
)

# ===================================================================
# 7. PERMISSION MODEL
# ===================================================================
add_heading("Permission Model", 2)
add_para(
    "Permissions are granted to groups, never to individual users. Production write "
    "access is held only by service principals: no human alters production state "
    "directly. The matrix below maps the 29 FTE in the Risk Division to their catalog "
    "access across the three environments."
)
add_table(
    headers=["Group", "FTE", "Catalog Scope", "Privileges", "Purpose"],
    rows_data=[
        [
            "retail_modelling-engineers",
            "10",
            "dev_retail_modelling, stg_retail_modelling, prod_retail_modelling",
            "MODIFY + CREATE on dev; SELECT on stg + prod",
            "Retail Modelling team builds and maintains scorecards (MVP starts here).",
        ],
        [
            "retail_insights-engineers",
            "9",
            "dev_retail_insights, stg_retail_insights, prod_retail_insights",
            "MODIFY + CREATE on dev; SELECT on stg + prod",
            "Retail Insights and Reporting team builds reporting outputs.",
        ],
        [
            "corporate_analytics-engineers",
            "6",
            "dev_corporate_analytics, stg_corporate_analytics, prod_corporate_analytics",
            "MODIFY + CREATE on dev; SELECT on stg + prod",
            "Corporate Analytics team builds and maintains its models.",
        ],
        [
            "financial_risk-engineers",
            "2",
            "dev_financial_risk, stg_financial_risk, prod_financial_risk",
            "MODIFY + CREATE on dev; SELECT on stg + prod",
            "Market Risk and Climate Risk modelling.",
        ],
        [
            "model-validation",
            "2",
            "All prod_<team> catalogs",
            "USE CATALOG, BROWSE, SELECT on monitoring + model registry schemas",
            "Model Risk and Assurance: validates models before promotion to champion.",
        ],
        [
            "<team>-cicd-sp (service principal)",
            "-",
            "stg_<team>",
            "USE CATALOG, MODIFY, CREATE",
            "CI service principal - runs integration tests in staging.",
        ],
        [
            "<team>-cd-sp (service principal)",
            "-",
            "prod_<team>",
            "USE CATALOG, MODIFY, CREATE",
            "Production deployment service principal - sole writer to prod.",
        ],
        [
            "platform-admins",
            "-",
            "All catalogs",
            "OWNERSHIP, MANAGE",
            "Platform team for ownership transfers and break-glass intervention.",
        ],
    ],
)
add_para(
    "Total: 29 FTE across four modelling teams plus the Model Risk and Assurance "
    "team. Service principals execute every production write; no individual user can "
    "alter production state directly."
)

doc.save(str(SRC))
print(f"Saved -> {SRC.name}")
